"""
============================================================
StudyAI Backend — File Parser Service
============================================================
Extracts raw text from uploaded files.

Supported formats:
  - PDF  (.pdf)   — via PyPDF2
  - DOCX (.docx)  — via python-docx
  - TXT  (.txt)   — plain read
  - MD   (.md)    — plain read (markdown as-is)
  - Raw text      — passed directly

Also provides word count, title inference, and text cleaning.
"""

from __future__ import annotations

import re
from pathlib import Path

from utils.logger import get_logger

logger = get_logger(__name__)

# ── Optional imports — degrade gracefully ─────────────────
try:
    import PyPDF2
    _HAS_PYPDF2 = True
except ImportError:
    _HAS_PYPDF2 = False
    logger.warning("PyPDF2 not installed — PDF parsing unavailable.")

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.warning("python-docx not installed — DOCX parsing unavailable.")


# ── Public API ─────────────────────────────────────────────

def parse_file(filepath: str) -> dict:
    """
    Extract text and metadata from a file.

    Args:
        filepath: Absolute path to the uploaded file.

    Returns:
        Dict with keys:
          ``text``       — full extracted text (str)
          ``word_count`` — number of words (int)
          ``char_count`` — number of characters (int)
          ``title``      — inferred title (str)
          ``extension``  — file extension without dot (str)
          ``page_count`` — pages for PDF, 1 otherwise (int)

    Raises:
        FileNotFoundError: If *filepath* does not exist.
        ValueError:        If the file extension is unsupported.
    """
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = path.suffix.lower().lstrip(".")
    logger.info("Parsing file: %s (type: %s)", path.name, ext)

    extractors = {
        "pdf":  _parse_pdf,
        "docx": _parse_docx,
        "txt":  _parse_plaintext,
        "md":   _parse_plaintext,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file type: '.{ext}'")

    result = extractor(path)
    result["extension"] = ext
    result["text"] = _clean_text(result.get("text", ""))
    result["word_count"] = _word_count(result["text"])
    result["char_count"] = len(result["text"])
    result["title"] = result.get("title") or _infer_title(path.name, result["text"])

    logger.info(
        "Parsed '%s' — %d words, %d chars, %d page(s)",
        path.name,
        result["word_count"],
        result["char_count"],
        result.get("page_count", 1),
    )
    return result


def parse_raw_text(text: str, title: str = "Untitled") -> dict:
    """
    Process raw text submitted directly (no file).

    Args:
        text:  The raw text content.
        title: Optional title for the material.

    Returns:
        Same structure as :func:`parse_file`.
    """
    cleaned = _clean_text(text)
    return {
        "text": cleaned,
        "word_count": _word_count(cleaned),
        "char_count": len(cleaned),
        "title": title,
        "extension": "txt",
        "page_count": 1,
    }


# ── Private Extractors ────────────────────────────────────

def _parse_pdf(path: Path) -> dict:
    """Extract text from a PDF file using PyPDF2."""
    if not _HAS_PYPDF2:
        raise ImportError("PyPDF2 is not installed. Run: pip install PyPDF2")

    pages: list[str] = []
    with open(path, "rb") as f:
        try:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
        except Exception as exc:
            logger.error("PDF read error for %s: %s", path.name, exc)
            raise ValueError(f"Could not read PDF: {exc}") from exc

    return {
        "text": "\n\n".join(pages),
        "page_count": page_count,
        "title": None,
    }


def _parse_docx(path: Path) -> dict:
    """Extract text from a DOCX file using python-docx."""
    if not _HAS_DOCX:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        raise ValueError(f"Could not open DOCX: {exc}") from exc

    # Extract title from core properties if available
    title = None
    try:
        title = doc.core_properties.title or None
    except Exception:
        pass

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return {
        "text": "\n".join(paragraphs),
        "page_count": 1,
        "title": title,
    }


def _parse_plaintext(path: Path) -> dict:
    """Read a plain-text or markdown file."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
    except OSError as exc:
        raise ValueError(f"Could not read file: {exc}") from exc

    return {
        "text": text,
        "page_count": 1,
        "title": None,
    }


# ── Text Utilities ────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Clean extracted text for LLM consumption.

    - Collapse multiple blank lines into a single blank line
    - Remove null bytes and form-feed characters
    - Strip leading/trailing whitespace
    """
    if not text:
        return ""
    text = text.replace("\x00", "").replace("\x0c", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _word_count(text: str) -> int:
    """Return the number of words in *text*."""
    return len(text.split()) if text else 0


def _infer_title(filename: str, text: str) -> str:
    """
    Attempt to infer a document title.

    Priority:
    1. First non-empty line of text (if short enough)
    2. Filename without extension
    """
    # Try first meaningful line
    for line in text.splitlines():
        line = line.strip()
        if line and len(line) <= 120:
            return line
    # Fallback: filename stem
    return Path(filename).stem.replace("_", " ").replace("-", " ").title()


def get_text_preview(text: str, max_chars: int = 500) -> str:
    """Return the first *max_chars* characters of *text* as a preview."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."
